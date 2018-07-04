from PyPDF2 import PdfFileWriter, PdfFileReader
import re
from docx import Document
from openpyxl import load_workbook
from docx.shared import Pt
from docx.shared import Cm
import os
import win32com.client


def get_fist_page_of_pdf(path):
    inputpdf = PdfFileReader(path)
    os.remove(path)
    output = PdfFileWriter()
    output.addPage(inputpdf.getPage(0))
    with open(path, "wb") as outputStream:
        output.write(outputStream)


def get_arabic_name(chinese_name, translatelist1):
    if chinese_name in translatelist1:
        arabic_name = translatelist1[chinese_name]
        return arabic_name
    else:
        print(chinese_name)
        print(chinese_name+" 没有阿拉伯语翻译")
        # add_Chinese_name_to_dictionary()
        exit()


def build_dictionary():
    workbook = load_workbook(filename='dictionary_csv/dictionary.xlsx', read_only=True)
    worksheet = workbook['Sheet1']
    translatelist1 = {}
    for row in worksheet.rows:
        translatelist1[row[0].value] = row[1].value
    return translatelist1


def build_chinese_name_to_ingredients_dictionary():
    workbook = load_workbook(filename='dictionary_csv/dictionary.xlsx', read_only=True)
    worksheet = workbook['Sheet2']
    chinese_name_to_ingredients_dictionary = {}
    chinese_name_to_row_number_dictionary = {}
    for row in worksheet.rows:
        chinese_name_to_ingredients_dictionary[row[0].value] = row[1].value
        chinese_name_to_row_number_dictionary[row[0].value] = row[6].value
    return chinese_name_to_ingredients_dictionary, chinese_name_to_row_number_dictionary


def write2docx(Ingredients_length_in_one_row, Ingredients_length_in_one_col, Chinese_Name,Ingredients_name_list, product_weight_number, product_weight, product_time, expire_time, Chinese_ProductPlace):
    stringlength = 80
    fontsize = 10
    margins_size = 0.25
    margins_size_top = margins_size
    margins_size_bot = margins_size
    margins_size_left = margins_size
    margins_size_right = margins_size
    row_num = Ingredients_length_in_one_col
    col_num = Ingredients_length_in_one_row
    translatelist1 = build_dictionary()
    Arabic_Name = get_arabic_name(Chinese_Name, translatelist1)
    Arabic_ProductPlace = get_arabic_name(Chinese_ProductPlace, translatelist1)
    Arabic_product_weight = get_arabic_name(product_weight, translatelist1)
    Arabic_Ingredients_name_list = [get_arabic_name(Ingredients_name_list[n], translatelist1) for n in
                                    range(len(Ingredients_name_list))]

    The_name = 'الاسم: '  # 商品名
    The_Ingredients_name = 'المكونات: '  # 配方
    The_serving_size = 'حجم الحصة: '  # serving size
    Product_date = 'تاريخ الإنتاج: '  # 生产日期
    expire_date = 'تاريخ انتهاء الصلاحية: '  # 过期日期
    Output_Product_name = The_name + Arabic_Name
    Output_Product_Ingredients_name = The_Ingredients_name
    for n in range(len(Arabic_Ingredients_name_list)):
        Output_Product_Ingredients_name = Output_Product_Ingredients_name + Arabic_Ingredients_name_list[n]
        if n != len(Arabic_Ingredients_name_list) - 1:
            Output_Product_Ingredients_name = Output_Product_Ingredients_name + '، '
    Our_Product_serving_size = The_serving_size + product_weight_number + ' ' + Arabic_product_weight
    Output_Production_date = Product_date + product_time + ' '
    Output_Expire_date = expire_date + expire_time + ' '
    Out_Product_place = Arabic_ProductPlace
    print(Chinese_Name.rjust(stringlength, ' '))
    print(Output_Product_name.ljust(stringlength, ' '))
    print(Output_Product_Ingredients_name.ljust(stringlength, ' '))
    print(Our_Product_serving_size.ljust(stringlength, ' '))
    print(Output_Production_date.ljust(stringlength, ' '))
    print(Output_Expire_date.ljust(stringlength, ' '))
    print(Out_Product_place.ljust(stringlength, ' '))

    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(margins_size_top)
        section.bottom_margin = Cm(margins_size_bot)
        section.left_margin = Cm(margins_size_left)
        section.right_margin = Cm(margins_size_right)
    table = document.add_table(rows=row_num, cols=col_num)
    for row_num_index in range(row_num):
        row_cells = table.rows[row_num_index].cells
        for col_num_index in range(col_num):
            p = row_cells[col_num_index].paragraphs[0]
            p.add_run(Chinese_Name + '\n').font.size = Pt(fontsize)
            p.add_run(Output_Product_name + '\n').font.size = Pt(fontsize)
            p.add_run(Output_Product_Ingredients_name + '\n').font.size = Pt(fontsize)
            p.add_run(Our_Product_serving_size + '\n').font.size = Pt(fontsize)
            p.add_run(Output_Production_date + '\n').font.size = Pt(fontsize)
            p.add_run(Output_Expire_date + '\n').font.size = Pt(fontsize)
            p.add_run(Out_Product_place).font.size = Pt(fontsize)
            p.alignment = 2
    document.save('temp_word/'+Chinese_Name+'-'+product_weight_number+product_weight + '.docx')
    #cannot get the number of page of the document Reference: https://stackoverflow.com/questions/24889845/number-of-pages-in-word-document
    #https: // stackoverflow.com / questions / 36193159 / page - number - python - docx
    wdFormatPDF = 17
    in_file = os.path.abspath('temp_word/'+Chinese_Name+'-'+product_weight_number+product_weight + '.docx')
    out_file = os.path.abspath('Out_put/'+Chinese_Name+'-'+product_weight_number+product_weight + '.pdf')
    word = win32com.client.Dispatch('Word.Application')
    #word = comtypes.client.CreateObject('Word.Application') 这个不好使了
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    get_fist_page_of_pdf(out_file)


def add_Ingredients_to_dictionary(Chinese_Name,list):
    wb = load_workbook(filename='dictionary_csv/dictionary.xlsx')
    old_sheet = wb.get_sheet_by_name('Sheet2')
    rows=(Chinese_Name,list,None,None,None,None,9)
    old_sheet.append(rows)
    wb.save(filename='dictionary_csv/dictionary.xlsx')


def get_Ingredients_name_list(list,Chinese_Name,Chinese_name_to_Ingredients_dictionary):
    if Chinese_Name in Chinese_name_to_Ingredients_dictionary:
        if list == None:
            list = Chinese_name_to_Ingredients_dictionary[Chinese_Name]
        Ingredients_name_list = re.split('，|,', list)
    else:
        if list!=None:
            add_Ingredients_to_dictionary(Chinese_Name,list)
            Ingredients_name_list = re.split('，|,', list)
        else:
            print(Chinese_Name+" 需要输入配方")
            exit()
    return Ingredients_name_list


def get_row_number(rownumber,Chinese_name_to_row_number_dictionary):
    if rownumber == None:
        rownumber = Chinese_name_to_row_number_dictionary[Chinese_Name]
    if rownumber == None:
        rownumber = 9
    return rownumber


def get_product_weight_number(list):
    product_weight_number_list = re.split('g|kg|ml|L|l|KG', list)
    return product_weight_number_list[0]


def get_product_weight(product_weight_number,list):
    product_weight_list = re.split(product_weight_number, list)
    return product_weight_list[1]

wb = load_workbook(filename='input_excel/input.xlsx', read_only=True)
ws = wb['Sheet1']
Chinese_name_to_Ingredients_dictionary,Chinese_name_to_row_number_dictionary=build_chinese_name_to_ingredients_dictionary()
jud=1
for col in ws.rows:
    if jud==1:
        jud=2
        continue
    if col[0].value==None:
        break
    else:
        Chinese_Name =col[0].value
        Ingredients_name_list = get_Ingredients_name_list(col[1].value,Chinese_Name,Chinese_name_to_Ingredients_dictionary)
        product_weight_number = get_product_weight_number(col[2].value)
        product_weight =  get_product_weight(product_weight_number,col[2].value)
        product_time = col[3].value.date().strftime("%Y/%m/%d")
        expire_time =col[4].value.date().strftime("%Y/%m/%d")
        Chinese_ProductPlace = col[5].value
        print(Ingredients_name_list)
        print(product_weight_number+product_weight)
        print(product_time)
        print(expire_time)
        print(Chinese_ProductPlace)
        rownumber=get_row_number(col[6].value,Chinese_name_to_row_number_dictionary)
        write2docx(5,rownumber, Chinese_Name, Ingredients_name_list, product_weight_number,
                   product_weight, product_time, expire_time, Chinese_ProductPlace)
